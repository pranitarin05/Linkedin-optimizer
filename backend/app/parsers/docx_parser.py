from docx import Document
import io
from typing import Dict, Any


def parse_docx(content: bytes) -> Dict[str, Any]:
    doc = Document(io.BytesIO(content))

    text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

    # Extract tables
    tables = []
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            table_data.append(row_data)
        tables.append(table_data)

    # Parse text into structured sections
    sections = _parse_text_sections(text)

    return {
        "raw_text": text,
        "sections": sections,
        "tables": tables,
        "confidence": _calculate_confidence(sections),
    }


def _parse_text_sections(text: str) -> Dict[str, str]:
    sections = {}
    lines = text.split("\n")

    current_section = "general"
    current_content = []

    section_keywords = {
        "experience": ["experience", "work history", "employment"],
        "education": ["education", "academic", "university", "college"],
        "skills": ["skills", "competencies", "proficiencies"],
        "certifications": ["certifications", "licenses", "credentials"],
        "summary": ["summary", "profile", "about", "objective"],
    }

    for line in lines:
        line_lower = line.lower().strip()
        found_section = False

        for section_name, keywords in section_keywords.items():
            if any(kw in line_lower for kw in keywords):
                if current_content:
                    sections[current_section] = "\n".join(current_content).strip()
                current_section = section_name
                current_content = []
                found_section = True
                break

        if not found_section:
            current_content.append(line)

    if current_content:
        sections[current_section] = "\n".join(current_content).strip()

    return sections


def _calculate_confidence(sections: Dict[str, str]) -> float:
    if not sections:
        return 0.0

    score = 0.0
    expected = ["experience", "education", "skills"]

    for section in expected:
        if section in sections and len(sections[section]) > 10:
            score += 0.33

    return min(score, 1.0)
